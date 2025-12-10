import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
import pandas as pd
from docx import Document
import os
import mammoth

def convert_xls_to_xlsx(file_path: str) -> str:
    """Converts .xls file to .xlsx using pandas with proper engine."""
    try:
        new_path = file_path + "x"
        # Use openpyxl engine to ensure proper xlsx format
        df = pd.read_excel(file_path, engine='xlrd')
        # Write with openpyxl to ensure compatibility with luckyexcel
        with pd.ExcelWriter(new_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        return new_path
    except Exception as e:
        print(f"Conversion failed: {e}")
        return file_path

def apply_excel_style(file_path: str, sheet_name: str, target_range: str, 
                      bold: bool = None, italic: bool = None, color: str = None, bg_color: str = None):
    """
    Applies styles to a range of cells in an Excel file.
    target_range can be a single cell 'A1', a range 'A1:B2', or a column 'A', or '1' for row.
    color and bg_color should be hex codes (e.g., 'FF0000').
    """
    try:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        
        # Determine cells to style
        cells_to_style = []
        
        if ":" in target_range:
             # Range like "A1:B10" or "A:A" (full column functionality limited in openpyxl iteration, parsing needed)
             # Openpyxl handles slices well: ws['A1':'C2']
             try:
                 rows = ws[target_range]
                 # If it's a single cell, rows might not be iterable as tuple of tuples the same way if strictly one cell? 
                 # actually ws['A1:A1'] returns ((Cell,),)
                 # ws['A:A'] returns cell column tuple
                 
                 # Flatten the tuple of tuples
                 if isinstance(rows, tuple):
                     for row in rows:
                         if isinstance(row, tuple):
                             cells_to_style.extend(row)
                         else:
                             cells_to_style.append(row)
                 else:
                     # single cell object
                     cells_to_style.append(rows)
             except Exception:
                 return f"Error: Invalid range format {target_range}"
        else:
            # Single cell 'A1' or Column 'A' or Row '1'
            if target_range.isalpha():
                # It's a column like 'A'
                # Iterate over existing cells in that column
                try:
                    cells_to_style = [cell for cell in ws[target_range]]
                except:
                    # Fallback for full column access if ws[col] fails
                    return f"Error: parsing column {target_range}"
            elif target_range.isdigit():
                 # It's a row like '1'
                try:
                    cells_to_style = [cell for cell in ws[int(target_range)]]
                except:
                    return f"Error: parsing row {target_range}"
            else:
                # Assume single cell 'A1'
                try:
                    cells_to_style = [ws[target_range]]
                except:
                    return f"Error: parsing cell {target_range}"

        # Apply styles
        for cell in cells_to_style:
            # Font updates
            current_font = cell.font
            new_bold = bold if bold is not None else current_font.bold
            new_italic = italic if italic is not None else current_font.italic
            new_color = color if color else current_font.color
            
            cell.font = Font(bold=new_bold, italic=new_italic, color=new_color, 
                             name=current_font.name, size=current_font.size)
            
        if bg_color:
            # Remove '#' if present
            hex_color = bg_color.lstrip('#')
            cell.fill = PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")

        # Atomic save
        tmp_path = file_path + ".tmp"
        wb.save(tmp_path)
        os.replace(tmp_path, file_path)
        return f"Applied styles to {target_range} successfully."
    except Exception as e:
        if os.path.exists(file_path + ".tmp"):
            os.remove(file_path + ".tmp")
        return f"Error applying styles: {e}"

def merge_excel_cells(file_path: str, sheet_name: str, range_string: str):
    """Merges cells in the specified range (e.g., 'A1:B2')."""
    try:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        ws.merge_cells(range_string)
        
        tmp_path = file_path + ".tmp"
        wb.save(tmp_path)
        os.replace(tmp_path, file_path)
        return f"Merged cells {range_string} successfully."
    except Exception as e:
        if os.path.exists(file_path + ".tmp"):
            os.remove(file_path + ".tmp")
        return f"Error merging cells: {e}"

def unmerge_excel_cells(file_path: str, sheet_name: str, range_string: str):
    """Unmerges cells in the specified range (e.g., 'A1:B2')."""
    try:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        ws.unmerge_cells(range_string)
        
        tmp_path = file_path + ".tmp"
        wb.save(tmp_path)
        os.replace(tmp_path, file_path)
        return f"Unmerged cells {range_string} successfully."
    except Exception as e:
        if os.path.exists(file_path + ".tmp"):
            os.remove(file_path + ".tmp")
        return f"Error unmerging cells: {e}"

def delete_excel_row(file_path: str, sheet_name: str, row_idx: int):
    """Deletes a row from an Excel sheet. row_idx is 1-based index."""
    try:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        ws.delete_rows(row_idx)
        
        tmp_path = file_path + ".tmp"
        wb.save(tmp_path)
        os.replace(tmp_path, file_path)
        return f"Row {row_idx} deleted successfully."
    except Exception as e:
        if os.path.exists(file_path + ".tmp"):
            os.remove(file_path + ".tmp")
        return f"Error deleting row: {e}"

def delete_excel_column(file_path: str, sheet_name: str, col_idx: str):
    """Deletes a column from an Excel sheet. col_idx can be a letter ('A') or 1-based index."""
    try:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        
        idx = col_idx
        if isinstance(col_idx, str) and col_idx.isalpha():
            # Convert letter to index
            from openpyxl.utils import column_index_from_string
            idx = column_index_from_string(col_idx)
        else:
             idx = int(col_idx)
            
        ws.delete_cols(idx)
        
        tmp_path = file_path + ".tmp"
        wb.save(tmp_path)
        os.replace(tmp_path, file_path)
        return f"Column {col_idx} deleted successfully."
    except Exception as e:
        if os.path.exists(file_path + ".tmp"):
             os.remove(file_path + ".tmp")
        return f"Error deleting column: {e}"

def get_preview_content(file_path: str) -> str:
    """Generates HTML preview for a file."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in ['.xlsx', '.xls']:
        try:
            df = pd.read_excel(file_path)
            return df.to_html(classes='min-w-full divide-y divide-gray-200', index=False)
        except Exception as e:
            return f"<div class='text-red-500'>Error reading Excel: {e}</div>"
            
    elif ext in ['.docx', '.doc']:
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                return f"<div class='prose max-w-none'>{result.value}</div>"
        except Exception as e:
            return f"<div class='text-red-500'>Error reading Word: {e}</div>"
            
    return "<div>Unsupported file type</div>"

def read_excel_structure(file_path: str) -> str:
    """Reads the sheet names and columns of an Excel file."""
    try:
        wb = openpyxl.load_workbook(file_path)
        info = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            # Get headers (assuming first row)
            headers = [cell.value for cell in ws[1] if cell.value]
            info.append(f"Sheet: {sheet}, Columns: {headers}")
        return "\n".join(info)
    except Exception as e:
        return f"Error: {e}"

def read_excel_values(file_path: str, sheet_name: str, range_string: str = None):
    """
    Reads values from a specific sheet. 
    range_string can be 'A1', 'A1:B2', or None (reads used range).
    """
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        
        data = []
        if range_string:
            # Check if it's a single cell
            if ":" not in range_string:
                 try:
                     val = ws[range_string].value
                     return str(val)
                 except:
                     pass 
            
            # Range
            try:
                rows = ws[range_string]
                if isinstance(rows, tuple):
                     for row in rows:
                        if isinstance(row, tuple):
                            data.append([str(cell.value) for cell in row])
                        else:
                            data.append([str(row.value)])
                else:
                    # Single cell object falling through
                    data.append([str(rows.value)])
            except Exception as e:
                return f"Error parsing range: {e}"
        else:
            # Read all rows
            for row in ws.iter_rows(values_only=True):
                data.append([str(c) for c in row])
                
        return str(data)
    except Exception as e:
        return f"Error reading values: {e}"

def add_excel_row(file_path: str, sheet_name: str, data: list):
    """Appends a row to an Excel sheet."""
    try:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        ws.append(data)
        
        tmp_path = file_path + ".tmp"
        wb.save(tmp_path)
        os.replace(tmp_path, file_path)
        return "Row added successfully."
    except Exception as e:
        if os.path.exists(file_path + ".tmp"):
            os.remove(file_path + ".tmp")
        return f"Error: {e}"

def write_excel_cell(file_path: str, sheet_name: str, cell: str, value: str):
    """Writes a value to a specific cell (e.g., 'A1')."""
    try:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name not in wb.sheetnames:
            return f"Error: Sheet {sheet_name} not found."
        ws = wb[sheet_name]
        ws[cell] = value
        
        tmp_path = file_path + ".tmp"
        wb.save(tmp_path)
        os.replace(tmp_path, file_path)
        return f"Wrote '{value}' to {cell} successfully."
    except Exception as e:
        if os.path.exists(file_path + ".tmp"):
            os.remove(file_path + ".tmp")
        return f"Error writing to cell: {e}"

def read_word_text(file_path: str) -> str:
    """Reads text from a Word file."""
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"Error: {e}"

def append_word_text(file_path: str, text: str):
    """Appends text to a Word file."""
    try:
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(file_path)
        return "Text appended successfully."
    except Exception as e:
        return f"Error: {e}"

def replace_word_text(file_path: str, old_text: str, new_text: str):
    """Replaces text in a Word file."""
    try:
        doc = Document(file_path)
        replaced_count = 0
        for p in doc.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
                replaced_count += 1
        
        # Also check tables if needed, but keeping it simple for paragraphs now
        
        doc.save(file_path)
        return f"Replaced {replaced_count} occurrences."
    except Exception as e:
        return f"Error: {e}"
